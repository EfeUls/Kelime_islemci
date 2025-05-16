import sys
from PyQt5 import uic
from PyQt5.QtWidgets import QApplication, QWidget, QMessageBox, QLabel, QVBoxLayout, QDialog, QPushButton
from PyQt5.QtGui import QPixmap, QColor, QFont
from PyQt5.QtCore import Qt, QUrl, pyqtSlot
from PyQt5.QtWebEngineWidgets import QWebEngineView
import serial.tools.list_ports
import serial
import matplotlib.pyplot as plt
from matplotlib.backends.backend_qt5agg import FigureCanvasQTAgg as FigureCanvas
from matplotlib.figure import Figure
import threading
import time
from PyQt5.QtGui import QFontDatabase
from PyQt5.QtWidgets import QApplication, QMainWindow, QComboBox, QTextEdit, QVBoxLayout
from PyQt5.uic import loadUi
from PyQt5.QtWidgets import QFileDialog
from PyQt5.QtWidgets import QPlainTextEdit
from PyQt5.QtGui import QTextCharFormat
from docx import Document
from PyQt5.QtGui import QTextCursor
import os
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
UI_EKRAN = os.path.join(BASE_DIR, "window", "kelime_islemci.ui")
class Pencere(QWidget):
    def __init__(self):
        super().__init__()
        uic.loadUi(UI_EKRAN , self)
        self.yaziboyutuBox.currentIndexChanged.connect(self.ayarlaYaziBoyutu)
        self.bold_button.clicked.connect(self.make_bold)
        self.italic_button.clicked.connect(self.make_italic)
        self.normal_button.clicked.connect(self.make_normal)
        self.cizgili_button.clicked.connect(self.make_cizgili)
        renkler = {"Siyah": "black", "Kirmizi": "red", "Mavi": "blue", "Yeşil": "green", "Sari": "yellow", "Beyaz": "white"}
        for isim, kod in renkler.items():
            self.colorBox.addItem(isim, kod)
        self.colorBox.currentIndexChanged.connect(self.renk_degistir)
        self.renk_degistir(self.colorBox.currentIndex())
        self.newButton.clicked.connect(self.yeni_dosya)

        self.openButton.clicked.connect(self.dosya_ac)
        self.saveButton.clicked.connect(self.kaydet)
        self.closeButton.clicked.connect(self.kapat)



    @pyqtSlot(int)
    def renk_degistir(self, index):
        secilen_renk_kodu = self.colorBox.itemData(index)
        if not secilen_renk_kodu:
            secilen_renk_kodu = self.colorBox.itemText(index).lower()
    
        cursor = self.metinAlani.textCursor()
        fmt = cursor.charFormat()
        fmt.setForeground(QColor(secilen_renk_kodu))
        cursor.mergeCharFormat(fmt)
        self.metinAlani.mergeCurrentCharFormat(fmt)
    
    def ayarlaYaziBoyutu(self, index):
            secilen_boyut_str = self.yaziboyutuBox.itemText(index)
            try:
                secilen_boyut = int(secilen_boyut_str)
                current_font = self.metinAlani.font()
                current_font.setPointSize(secilen_boyut)
                self.metinAlani.setFont(current_font)
            except ValueError:
                print("Geçersiz yazı boyutu değeri!")
    def make_bold(self):
        cursor = self.metinAlani.textCursor()
        fmt = cursor.charFormat()
        fmt.setFontWeight(QFont.Bold)
        cursor.mergeCharFormat(fmt)
        self.metinAlani.mergeCurrentCharFormat(fmt)



    def make_italic(self):
        cursor = self.metinAlani.textCursor()
        fmt = cursor.charFormat()
        fmt.setFontItalic(True)
        cursor.mergeCharFormat(fmt)
        self.metinAlani.mergeCurrentCharFormat(fmt)


    def make_normal(self):
        cursor = self.metinAlani.textCursor()
        fmt = QTextCharFormat()
        fmt.setFontWeight(QFont.Normal)
        fmt.setFontItalic(False)
        fmt.setFontUnderline(False)
    
        if cursor.hasSelection():
            cursor.mergeCharFormat(fmt)
        else:
            self.metinAlani.setCurrentCharFormat(fmt)

    
    def make_cizgili(self):
        cursor = self.metinAlani.textCursor()
        fmt = cursor.charFormat()
        fmt.setFontUnderline(True)
        cursor.mergeCharFormat(fmt)
        self.metinAlani.mergeCurrentCharFormat(fmt)




    def kaydet(self):
        if not hasattr(self, 'metinAlani'):
           QMessageBox.warning(self, "Uyarı", "Metin alanı bulunamadı!")
           return

        dosya_adi, _ = QFileDialog.getSaveFileName(
            self, "Kaydet", "", "Word Dosyaları (*.docx);;Tüm Dosyalar (*)"
        )
    
        if not dosya_adi:
            return
    
        try:
            if isinstance(self.metinAlani, QPlainTextEdit):
                text_edit = QTextEdit()
                text_edit.setPlainText(self.metinAlani.toPlainText())
                html_content = text_edit.toHtml()
            else:
                html_content = self.metinAlani.toHtml()
    
            import tempfile
            with tempfile.NamedTemporaryFile(suffix=".html", delete=False) as temp_file:
                temp_file.write(html_content.encode('utf-8'))
                temp_path = temp_file.name
    
            document = Document()
            #document.add_paragraph("Biçimlendirilmiş Metin:")
            
            from bs4 import BeautifulSoup
            soup = BeautifulSoup(html_content, 'html.parser')
            
            for p in soup.find_all('p'):
                paragraph = document.add_paragraph()
                for content in p.contents:
                    if content.name == 'span' or content.name == 'font':
                        run = paragraph.add_run(content.text)
                        if 'style' in content.attrs:
                            if 'font-weight:bold' in content['style']:
                                run.bold = True
                            if 'font-style:italic' in content['style']:
                                run.italic = True
                    else:
                        paragraph.add_run(str(content))
    
            document.save(dosya_adi)
            QMessageBox.information(self, "Başarılı", f"Metin başarıyla kaydedildi: {dosya_adi}")
            return True
        except Exception as e:
            QMessageBox.critical(self, "Hata", f"Kaydetme sırasında bir hata oluştu: {e}")
            return False
        finally:
            import os
            if 'temp_path' in locals() and os.path.exists(temp_path):
                os.remove(temp_path)
    

    def kapat(self):
        if self.metinAlani and self.metinAlani.toPlainText(): 
            onay_penceresi = KaydetOnayPenceresi(self)
            result = onay_penceresi.exec_() 

            if result == QDialog.Accepted: 
                sys.exit()
            elif result == QDialog.Rejected:
                if self.kaydet(): 
                    sys.exit()
                
        else:
            sys.exit() 
        
    def dosya_ac(self):
        
        dosya_adi, _ = QFileDialog.getOpenFileName(
            self, 
            "Dosya Aç", 
            "",  
            "Metin Dosyaları (*.txt);;Word Dosyaları (*.docx);;Tüm Dosyalar (*)"
        )
        
        if not dosya_adi:  
            return
        
        try:
            if dosya_adi.endswith('.txt'):
                with open(dosya_adi, 'r', encoding='utf-8') as dosya:
                    icerik = dosya.read()
                    
                    if isinstance(self.metinAlani, QPlainTextEdit):
                        self.metinAlani.setPlainText(icerik)
                    elif isinstance(self.metinAlani, QTextEdit):
                        self.metinAlani.setPlainText(icerik) 
                        
            elif dosya_adi.endswith('.docx'):
                document = Document(dosya_adi)
                full_text = []
                for para in document.paragraphs:
                    full_text.append(para.text)
                
                icerik = '\n'.join(full_text)
                if isinstance(self.metinAlani, QPlainTextEdit):
                    self.metinAlani.setPlainText(icerik)
                elif isinstance(self.metinAlani, QTextEdit):
                    self.metinAlani.setPlainText(icerik)
                    
            else:
                QMessageBox.warning(self, "Uyarı", "Desteklenmeyen dosya formatı!")
                return
                
        except UnicodeDecodeError:
            QMessageBox.critical(self, "Hata", "Dosya okunurken kodlama hatası oluştu!")
        except PermissionError:
            QMessageBox.critical(self, "Hata", "Dosyaya erişim izni reddedildi!")
        except Exception as e:
            QMessageBox.critical(
                self, 
                "Hata", 
                f"Dosya açılırken beklenmeyen bir hata oluştu:\n{str(e)}"
            )
    def yeni_dosya(self):
        self.metinAlani.clear()


    


class KaydetOnayPenceresi(QDialog):
      def __init__(self , parent=None):
          super().__init__(parent)
          uic.loadUi("c:/Users/eulas/Desktop/PyQt5Ders/Visual-Programming-PyQT-Lecture-Notes-1/kapatt.ui", self)
          self.setWindowTitle("Kaydetme Onayı")
          #self.closeButton.clicked.connect(self.kapat)

          self.evet_button.clicked.connect(self.accept)
          self.hayir_button.clicked.connect(self.reject)
           

if __name__ == "__main__":
    sys = QApplication([])
    pencere = Pencere()
    pencere.show()
    sys.exec_()